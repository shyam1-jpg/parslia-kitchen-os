#!/usr/bin/env python3
"""Build a single polished Vedanta Kitchen combined risk assessment Word document."""

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

IMG_DIR = "/workspace/vedanta-risk-assessments/images"
BRAND_DIR = f"{IMG_DIR}/brand"
PRODUCT_DIR = f"{IMG_DIR}/products"
STEPS_DIR = f"{IMG_DIR}/steps"

# Brand colours — The Vedanta Way (thevedanta.org)
NAVY = RGBColor(0x1A, 0x1A, 0x1A)  # charcoal / near-black
TEAL = RGBColor(0x8A, 0x73, 0x4A)  # soft gold accent
DARK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_ROW = "F5F2EC"
HEADER_BG = "1A1A1A"
ALT_HEADER = "8A734A"


def set_run(run, *, size=11, bold=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def write_cell(cell, text, *, bold=False, color=DARK, size=10, center=False, fill=None):
    clear_paragraph(cell.paragraphs[0])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=NAVY, font="Georgia")
    elif level == 2:
        set_run(run, size=13, bold=True, color=TEAL, font="Georgia")
    else:
        set_run(run, size=11, bold=True, color=NAVY)
    return p


def add_body(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    clear_paragraph(p)
    run = p.add_run(text)
    set_run(run, size=11, color=DARK)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. {item}")
        set_run(run, size=11, color=DARK)


def info_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, (label, value) in enumerate(rows):
        write_cell(table.rows[i].cells[0], label, bold=True, size=10, fill=LIGHT_ROW, color=NAVY)
        write_cell(table.rows[i].cells[1], value, size=10)
        table.rows[i].cells[0].width = Cm(5.5)
        table.rows[i].cells[1].width = Cm(11)
    doc.add_paragraph()


def risk_table(doc, hazards):
    """hazards: list of (hazard, people, before, measures, after)"""
    table = doc.add_table(rows=1 + len(hazards), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Hazard", "People at Risk", "Risk Before", "Control Measures", "Risk After"]
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=9, center=True, fill=HEADER_BG)
    for r, row in enumerate(hazards, 1):
        fill = LIGHT_ROW if r % 2 == 0 else None
        for c, val in enumerate(row):
            write_cell(
                table.rows[r].cells[c],
                val,
                size=9,
                center=(c in (2, 4)),
                fill=fill,
                bold=(c in (2, 4)),
                color=TEAL if c in (2, 4) else DARK,
            )
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def add_picture(doc, filename, width_inches=6.3, *, folder=None):
    base = folder or IMG_DIR
    path = f"{base}/{filename}" if "/" not in filename else filename
    if folder:
        path = f"{folder}/{filename}"
    elif not filename.startswith("/"):
        # try root images, then brand, then products
        for candidate in (
            f"{IMG_DIR}/{filename}",
            f"{BRAND_DIR}/{filename}",
            f"{PRODUCT_DIR}/{filename}",
            f"{STEPS_DIR}/{filename}",
        ):
            if os.path.exists(candidate):
                path = candidate
                break
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    return p


def add_steps_guide(doc, filename, caption="Step-by-step visual guide"):
    add_heading_styled(doc, caption, 2)
    try:
        add_picture(doc, filename, width_inches=6.4)
    except Exception:
        add_body(doc, f"(Step guide image unavailable: {filename})", size=9)


def cover_page(doc):
    try:
        add_picture(doc, "Vedanta-Way-Ltd-text-banner.jpg", width_inches=6.5)
    except Exception:
        pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    run = title.add_run("THE VEDANTA WAY LIMITED")
    set_run(run, size=12, bold=True, color=TEAL, font="Georgia")

    main = doc.add_paragraph()
    main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main.paragraph_format.space_before = Pt(6)
    run = main.add_run("Kitchen Safety Brochure &\nCombined Risk Assessment Pack")
    set_run(run, size=24, bold=True, color=NAVY, font="Georgia")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(10)
    run = sub.add_run(
        "Equipment · Sylvester Keal Chemicals (original brochure products) ·\n"
        "Step-by-step Do / Don’t / Care staff training guides"
    )
    set_run(run, size=11, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    run = meta.add_run(
        "The Vedanta Kitchen & Retreat Centre\n"
        "Branston Hall, Lincoln Road, Branston, Lincoln LN4 1PD\n"
        "Company: The Vedanta Way Limited\n"
        "Website: https://thevedanta.org/\n\n"
        f"Assessor: Shyam Prasad\n"
        f"Document date: {date.today().strftime('%d %B %Y')}\n"
        f"Review due: 27 July 2027"
    )
    set_run(run, size=10, color=DARK)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run(
        "Includes:\n"
        "TRK70 · Rational Oven · SK Chemicals / COSHH · Thermomix · Caso ·\n"
        "Ninja · KitchenAid · Waring · Knives / Mandoline ·\n"
        "Lincat Fryer · Dough Mixer · Blast Chiller · Bratt Pan ·\n"
        "Wrapmaster Cutters · Buffalo Induction · Heavy Tray Storage ·\n"
        "Kitchen, Catering, FOH & Generic Assessments"
    )
    set_run(run, size=9, color=MUTED)
    page_break(doc)


def add_do_dont_care(doc, do_items, dont_items, care_items):
    add_heading_styled(doc, "Quick visual guide — Do / Don’t / Care", 2)
    add_body(doc, "DO", bold=True, space_after=2)
    for item in do_items:
        add_bullet(doc, item)
    add_body(doc, "DON’T", bold=True, space_after=2)
    for item in dont_items:
        add_bullet(doc, item)
    add_body(doc, "CARE", bold=True, space_after=2)
    for item in care_items:
        add_bullet(doc, item)


def contents_page(doc, sections):
    add_heading_styled(doc, "Contents", 1)
    add_body(doc, "Use this pack as the master risk assessment file for kitchen and catering operations.")
    for i, name in enumerate(sections, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{i}.  {name}")
        set_run(run, size=11, color=DARK)
    page_break(doc)


def equipment_section(
    doc,
    title,
    brand,
    model,
    description,
    procedure,
    hazards,
    ppe,
    *,
    extra_notes=None,
    steps_image=None,
):
    add_heading_styled(doc, title, 1)
    if steps_image:
        add_steps_guide(doc, steps_image, "Step-by-step safe use")
    add_heading_styled(doc, "1. Equipment Information", 2)
    info_table(
        doc,
        [
            ("Location", "The Vedanta Kitchen & Retreat Centre"),
            ("Assessor", "Shyam Prasad"),
            ("Brand", brand),
            ("Model", model),
            ("Date Completed", date.today().strftime("%d/%m/%Y")),
            ("Review Date", "04/11/2025"),
        ],
    )

    add_heading_styled(doc, "2. Description of Equipment", 2)
    add_body(doc, description)

    add_heading_styled(doc, "3. Operating Procedure (How to Use)", 2)
    add_numbered(doc, procedure)

    add_heading_styled(doc, "4. Risk Assessment", 2)
    risk_table(doc, hazards)

    add_heading_styled(doc, "5. Required PPE", 2)
    for item in ppe:
        add_bullet(doc, item)

    if extra_notes:
        add_heading_styled(doc, "6. Additional Notes", 2)
        for note in extra_notes:
            add_bullet(doc, note)

    page_break(doc)


def area_section(doc, title, scope, hazards, controls, ppe, notes=None):
    add_heading_styled(doc, title, 1)
    add_heading_styled(doc, "1. Scope", 2)
    add_body(doc, scope)

    add_heading_styled(doc, "2. Key Information", 2)
    info_table(
        doc,
        [
            ("Location", "The Vedanta Kitchen & Retreat Centre"),
            ("Assessor", "Shyam Prasad"),
            ("Assessment type", "Area / activity risk assessment"),
            ("Date Completed", date.today().strftime("%d/%m/%Y")),
            ("Review Date", "04/11/2025"),
        ],
    )

    add_heading_styled(doc, "3. Risk Assessment", 2)
    risk_table(doc, hazards)

    add_heading_styled(doc, "4. Control Summary", 2)
    for c in controls:
        add_bullet(doc, c)

    add_heading_styled(doc, "5. Required PPE", 2)
    for item in ppe:
        add_bullet(doc, item)

    if notes:
        add_heading_styled(doc, "6. Additional Notes", 2)
        for n in notes:
            add_bullet(doc, n)

    page_break(doc)


def detailed_risk_table(doc, hazards):
    """hazards: list of (hazard, harm, initial, controls, residual)"""
    table = doc.add_table(rows=1 + len(hazards), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Hazard", "Possible harm", "Initial risk", "Required control measures", "Residual risk"]
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=8, center=True, fill=HEADER_BG)
    for r, row in enumerate(hazards, 1):
        fill = LIGHT_ROW if r % 2 == 0 else None
        for c, val in enumerate(row):
            write_cell(
                table.rows[r].cells[c],
                val,
                size=8,
                center=(c in (2, 4)),
                fill=fill,
                bold=(c in (2, 4)),
                color=TEAL if c in (2, 4) else DARK,
            )
    doc.add_paragraph()


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=11, bold=True, color=NAVY)
    return p


def trk70_section(doc):
    add_heading_styled(
        doc,
        "Electrolux / Dito Sama TRK70 — Combined Cutter & Vegetable Slicer Risk Assessment",
        1,
    )
    try:
        add_picture(doc, "trk70-do-dont-care.png")
    except Exception:
        pass
    add_steps_guide(doc, "steps-trk70.png", "TRK70 — Step-by-step safe use")
    add_do_dont_care(
        doc,
        [
            "Use only the supplied food pusher",
            "Keep hopper / lid closed and interlocks engaged",
            "Wear cut-resistant gloves when changing blades",
            "Unplug before cleaning or clearing blockages",
        ],
        [
            "Never put hands or utensils in the feed chute while powered",
            "Never run the machine with a guard open or bypassed",
            "Never leave blades in sink water or loose in drawers",
            "Never allow untrained staff to operate the TRK70",
        ],
        [
            "Inspect guards, cable and blades before every use",
            "Store discs in the protective rack",
            "Clean and sanitise food-contact parts after use",
            "Label and remove faulty equipment from service",
        ],
    )

    add_heading_styled(doc, "1. Workplace & Equipment Information", 2)
    info_table(
        doc,
        [
            ("Workplace", "The Vedanta Kitchen & Retreat Centre"),
            ("Department", "Main Kitchen"),
            ("Equipment", "Electrolux / Dito Sama TRK70 Combined Cutter and Vegetable Slicer"),
            ("Manufacturer / model", "Electrolux Professional / Dito Sama TRK70"),
            ("Persons completing the task", "Trained and authorised kitchen staff"),
            ("Assessor", "Shyam Prasad"),
            ("Assessment date", "27 July 2026"),
            ("Review date", "27 July 2027, or sooner after an accident, near miss, equipment change or safety concern"),
            ("Manager responsible", "[Name]"),
        ],
    )

    add_heading_styled(doc, "2. Activity Covered", 2)
    add_body(
        doc,
        "Operating, loading, unloading, changing cutting discs or blades, clearing blockages, "
        "dismantling, cleaning, reassembling and maintaining the TRK70.",
    )
    add_body(
        doc,
        "The TRK70 uses interchangeable rotating blades and discs for cutting, chopping, slicing, "
        "grating, shredding and dicing food. The machine incorporates a magnetic safety system and "
        "motor brake intended to stop operation when the hopper, vegetable preparation lever or "
        "cutter lid is not correctly closed. These safety systems must never be defeated or bypassed.",
    )

    add_heading_styled(doc, "3. Persons at Risk", 2)
    for person in [
        "Chefs and kitchen assistants",
        "Kitchen porters and cleaning staff",
        "Maintenance contractors",
        "Employees working close to the machine",
        "New, inexperienced or young workers",
        "Consumers, where poor cleaning or allergen controls could contaminate food",
    ]:
        add_bullet(doc, person)

    add_heading_styled(doc, "4. Risk-Rating System", 2)
    add_body(doc, "Likelihood", bold=True, space_after=2)
    for item in [
        "1 – Rare",
        "2 – Unlikely",
        "3 – Possible",
        "4 – Likely",
        "5 – Almost certain",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Severity", bold=True, space_after=2)
    for item in [
        "1 – Minor injury",
        "2 – Injury requiring first aid",
        "3 – Injury requiring medical treatment",
        "4 – Serious injury",
        "5 – Fatality or life-changing injury",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Risk score = Likelihood × Severity", bold=True, space_after=2)
    for item in [
        "1–4: Low",
        "5–9: Medium",
        "10–15: High",
        "16–25: Very high",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "5. Significant Hazards and Control Measures", 2)
    detailed_risk_table(
        doc,
        [
            (
                "Contact with rotating cutter blade or slicing disc",
                "Deep cuts, amputation or life-changing injury",
                "15 – High",
                "Only trained and authorised staff may operate. All guards, hopper, lid, pushers and interlocks must be correctly fitted. Hands, knives, spoons and utensils must never enter the feed opening or bowl while connected to power. Feed food only with the supplied pusher.",
                "5 – Medium",
            ),
            (
                "Reaching into the machine while it is running",
                "Severe cuts, crushing or amputation",
                "15 – High",
                "Never reach beneath the hopper, through the feed chute or into the cutter bowl while operating. Stop and isolate from the electrical supply before inspecting or touching internal components.",
                "5 – Medium",
            ),
            (
                "Unexpected start-up during assembly, cleaning or blockage removal",
                "Contact with moving blades resulting in serious injury",
                "15 – High",
                "Switch off using the main control, turn off the rear isolation switch and unplug before dismantling, cleaning, changing blades or clearing a blockage. Keep the plug under the operator’s control until the task is complete.",
                "5 – Medium",
            ),
            (
                "Defective, removed or bypassed guard / interlock",
                "Access to dangerous moving parts",
                "15 – High",
                "Check lid, hopper, pusher, guard and safety interlock before every use. Never bypass, tape down or modify a safety device. If the machine runs with a guard open, stop, unplug, label “DO NOT USE” and report to the manager.",
                "5 – Medium",
            ),
            (
                "Handling, fitting or removing sharp blades and discs",
                "Cuts and puncture wounds to hands and fingers",
                "12 – High",
                "Switch off and unplug before handling blades. Hold only by hub, handle or non-cutting edge. Wear cut-resistant gloves when changing, transporting or manually cleaning blades. Do not test sharpness with fingers.",
                "4 – Low",
            ),
            (
                "Blades stored incorrectly",
                "Cuts while reaching into drawers or cupboards",
                "12 – High",
                "Store each blade and disc in a designated rack, protective holder or labelled container. Cutting edges must not be left exposed. Never leave blades submerged in sinks or hidden beneath other equipment.",
                "4 – Low",
            ),
            (
                "Blocked machine or food becoming stuck",
                "Operator may attempt to push food with hands or utensils",
                "15 – High",
                "Stop and unplug before clearing any blockage. Wait until all movement has stopped. Remove hopper, lid or attachment per manufacturer instructions. Use an appropriate cleaning tool after isolation; never use hands while connected.",
                "5 – Medium",
            ),
            (
                "Food or broken blade components being ejected",
                "Cuts, bruising or eye injury",
                "12 – High",
                "Inspect blades, discs, bowl, shaft and attachments for cracks, distortion or damage before use. Secure attachments correctly. Do not process bones, frozen-solid food or unsuitable materials. Do not overload. Stand clear when starting.",
                "4 – Low",
            ),
            (
                "Loose clothing, jewellery or long hair becoming caught",
                "Entanglement, pulling or impact injury",
                "12 – High",
                "Wear close-fitting kitchen clothing. Tie back and secure long hair. Remove loose jewellery, scarves and lanyards. Secure sleeves. Do not distract anyone operating the machine.",
                "4 – Low",
            ),
            (
                "Electrical damage, wet plug or damaged cable",
                "Electric shock, burns or fire",
                "10 – High",
                "Visually inspect plug, cable, socket and machine before use. Keep connections dry. Do not use with wet hands or pull/lift by the cable. Damaged equipment must be unplugged, labelled and removed from service until repaired by a competent person.",
                "5 – Medium",
            ),
            (
                "Water entering the motor or electrical controls during cleaning",
                "Electric shock, equipment failure or fire",
                "10 – High",
                "Unplug before cleaning. Never hose down, immerse or pressure-wash the motor base. Clean with a damp cloth using the manufacturer-approved method. Removable food-contact parts may be cleaned separately. Ensure components are dry before reassembly.",
                "5 – Medium",
            ),
            (
                "Cleaning chemicals",
                "Skin irritation, eye damage or breathing difficulty",
                "6 – Medium",
                "Follow the relevant COSHH assessment and product instructions. Use correct dilution. Wear gloves and eye protection where specified. Never mix chemicals. Store away from food and rinse food-contact components thoroughly.",
                "3 – Low",
            ),
            (
                "Wet or contaminated floor around the machine",
                "Slips, falls, strains or impact injuries",
                "9 – Medium",
                "Position on a stable work surface. Clean spillages immediately. Display a wet-floor sign where necessary. Keep the power cable away from walkways. Wear slip-resistant safety footwear.",
                "3 – Low",
            ),
            (
                "Machine positioned on an unstable or unsuitable surface",
                "Machine movement, falling equipment or operator injury",
                "12 – High",
                "Place on a strong, level, dry and stable surface at a comfortable working height. Ensure the base is secure before operation. Do not operate close to the edge of a workbench.",
                "4 – Low",
            ),
            (
                "Lifting or moving the machine and attachments",
                "Back, shoulder, hand or foot injuries",
                "9 – Medium",
                "Avoid moving the complete machine unnecessarily. Assess weight before lifting. Use a trolley or two-person lift where required. Remove loose attachments before transport. Keep the route clear and use correct manual-handling techniques.",
                "4 – Low",
            ),
            (
                "Repetitive loading, pushing and awkward posture",
                "Musculoskeletal injury or fatigue",
                "6 – Medium",
                "Position the machine at a suitable height. Keep ingredients close to the operator. Avoid twisting or overreaching. Rotate tasks during prolonged production and take appropriate breaks.",
                "2 – Low",
            ),
            (
                "Excessive machine noise, vibration or overheating",
                "Hearing discomfort, equipment failure or fire",
                "6 – Medium",
                "Stop if unusual noise, vibration, smell, smoke or overheating occurs. Unplug and report the fault. Do not continue operating damaged equipment. Keep ventilation openings clear.",
                "2 – Low",
            ),
            (
                "Poor cleaning or food remaining inside attachments",
                "Bacterial contamination and food poisoning",
                "12 – High",
                "Dismantle and clean food-contact parts after every use and between incompatible foods. Inspect hidden areas, shafts, seals, lids and feed components. Sanitise using the approved kitchen procedure and dry before reassembly.",
                "4 – Low",
            ),
            (
                "Allergen cross-contamination",
                "Allergic reaction, potentially severe or fatal",
                "12 – High",
                "Follow the kitchen allergen-management procedure. Clean and sanitise bowl, blades, discs, hopper, pushers, utensils and surrounding work surface between allergen and non-allergen recipes. Use dedicated equipment where required and verify cleaning before production begins.",
                "4 – Low",
            ),
            (
                "Use by an untrained or unauthorised person",
                "Incorrect assembly, unsafe operation or serious injury",
                "15 – High",
                "Operators must receive practical training and be assessed as competent. New employees must be directly supervised. Cleaning staff must be trained in isolation and safe blade handling. Operating instructions and this assessment must remain accessible.",
                "5 – Medium",
            ),
            (
                "Maintenance or repairs by an unauthorised person",
                "Electric shock, unexpected movement or defective safety systems",
                "15 – High",
                "Operators may only perform normal cleaning and user checks. Electrical repairs, motor work, interlock adjustments and internal maintenance must be completed by an authorised competent engineer.",
                "5 – Medium",
            ),
        ],
    )

    add_heading_styled(doc, "6. Safe Operating Procedure", 2)

    add_subheading(doc, "Before use")
    add_numbered(
        doc,
        [
            "Confirm that the operator is trained and authorised.",
            "Check that the machine is clean, dry and correctly positioned.",
            "Inspect the plug, cable, controls, bowl, lid, hopper, pusher, blades and discs.",
            "Confirm that all guards and safety interlocks are present and undamaged.",
            "Select the correct blade, disc and speed for the food being processed.",
            "Fit the attachment while the machine is unplugged.",
            "Confirm that the blade or disc is properly secured.",
            "Secure long hair and remove loose jewellery.",
            "Keep the floor and surrounding work area clean and dry.",
            "Do not operate the machine if any component is missing, loose or damaged.",
        ],
    )

    add_subheading(doc, "During operation")
    add_numbered(
        doc,
        [
            "Close and secure the bowl lid or vegetable-slicer hopper.",
            "Start the machine only after checking that no person is touching the attachments.",
            "Use only the supplied pusher to feed food.",
            "Never place hands or utensils inside the feed opening or bowl.",
            "Do not overload the machine.",
            "Remain with the machine while it is operating.",
            "Stop immediately if there is unusual noise, vibration, smell, heat or movement.",
            "Do not remove any guard until the machine has stopped completely and has been unplugged.",
        ],
    )

    add_subheading(doc, "Blade or disc changes")
    add_numbered(
        doc,
        [
            "Switch the machine off.",
            "Turn off the rear isolation switch.",
            "Unplug the machine.",
            "Wait until all movement has stopped.",
            "Wear suitable cut-resistant gloves.",
            "Hold the blade or disc by its hub or non-cutting edge.",
            "Place the removed blade directly into its protective storage rack.",
            "Fit the replacement securely before reconnecting the machine.",
        ],
    )

    add_subheading(doc, "Clearing a blockage")
    add_numbered(
        doc,
        [
            "Press the stop control.",
            "Switch off the main isolation switch.",
            "Unplug the machine.",
            "Keep control of the plug.",
            "Wait for the blade or disc to stop completely.",
            "Carefully dismantle the appropriate attachment.",
            "Remove the blockage without placing fingers against cutting edges.",
            "Inspect for blade damage before reassembling.",
            "Do not restart the machine until all guards are correctly fitted.",
        ],
    )

    add_subheading(doc, "Cleaning")
    add_numbered(
        doc,
        [
            "Switch off and unplug the machine.",
            "Wear cut-resistant gloves when handling blades.",
            "Remove the blade or disc first and place it in a safe position.",
            "Never leave a blade hidden or submerged in washing-up water.",
            "Dismantle removable food-contact components.",
            "Wash, rinse, sanitise and dry components according to the kitchen cleaning procedure.",
            "Do not immerse or hose down the motor base.",
            "Inspect components for cracks, damage or food deposits.",
            "Reassemble only when all parts are clean and dry.",
            "Store blades and discs in their designated protective rack.",
        ],
    )

    add_heading_styled(doc, "7. Emergency Arrangements", 2)
    add_subheading(doc, "Serious cut or amputation")
    for item in [
        "Stop and isolate the machine immediately.",
        "Call 999 for a serious or uncontrolled injury.",
        "Apply firm pressure using a sterile dressing.",
        "Do not remove embedded objects.",
        "Preserve any severed part in a clean sealed bag, place that bag within another bag containing ice, and send it with the injured person.",
        "Report the accident immediately and secure the machine for investigation.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "Electric shock")
    for item in [
        "Do not touch the injured person until the electrical supply has been safely isolated.",
        "Switch off the supply or unplug the machine where safe.",
        "Call 999.",
        "Begin first aid or CPR if trained and instructed to do so.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "Equipment malfunction")
    for item in [
        "Press stop and disconnect the power.",
        "Label the machine “DO NOT USE – FAULTY.”",
        "Inform the kitchen manager.",
        "Arrange inspection by a competent engineer.",
        "Do not return the machine to service until formally declared safe.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "8. Training Requirements", 2)
    add_body(doc, "Operators must be trained in:")
    for item in [
        "Correct assembly and attachment selection",
        "Operation of the controls and isolation switch",
        "Use of guards, hopper and food pushers",
        "Safe blade and disc handling",
        "Blockage-clearing procedure",
        "Cleaning and sanitising",
        "Allergen cross-contamination controls",
        "Pre-use inspection and fault reporting",
        "Emergency procedures",
    ]:
        add_bullet(doc, item)
    add_body(
        doc,
        "Training must be recorded and refreshed following an accident, near miss, unsafe practice, "
        "equipment modification or significant period without using the machine.",
    )

    add_heading_styled(doc, "9. Inspection and Maintenance", 2)
    for item in [
        "Carry out a visual pre-use check before every shift.",
        "Inspect guards, lids, pushers and interlocks before every use.",
        "Record and report damage immediately.",
        "Maintain the machine according to the manufacturer’s instructions.",
        "Arrange electrical inspection and testing based on usage, environment and previous inspection results.",
        "Only competent authorised persons may repair or adjust the machine.",
        "Keep servicing, repair and inspection records.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "10. Residual Risk Statement", 2)
    add_body(
        doc,
        "Provided that all listed controls are implemented, operators are trained, guards and interlocks "
        "remain functional, and the machine is isolated before blade handling, cleaning or blockage removal, "
        "the remaining risk is considered medium to low and tolerable with continued supervision.",
    )
    add_body(doc, "The machine must not be used where:", bold=True, space_after=2)
    for item in [
        "A guard, lid, hopper or pusher is missing",
        "A safety interlock is defective",
        "The blade or disc is damaged",
        "The machine starts while a guard is open",
        "The cable, plug or controls are damaged",
        "There is unusual noise, vibration, heat, smoke or smell",
        "The operator has not received suitable training",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "11. Required PPE", 2)
    for item in [
        "Apron / chef whites",
        "Non-slip / slip-resistant footwear",
        "Cut-resistant gloves when changing, transporting or manually cleaning blades and discs",
        "Eye protection and chemical gloves where specified by COSHH for cleaning chemicals",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "12. Approval", 2)
    info_table(
        doc,
        [
            ("Assessor’s name", "Shyam Prasad"),
            ("Assessor signature", ""),
            ("Assessor date", "27 July 2026"),
            ("Manager’s name", ""),
            ("Manager signature", ""),
            ("Manager date", ""),
        ],
    )

    add_heading_styled(doc, "13. Employee Acknowledgement", 2)
    add_body(
        doc,
        "I confirm that I have read and understood this risk assessment, received suitable instruction "
        "and training, and agree to follow the stated control measures and safe operating procedure.",
    )
    table = doc.add_table(rows=6, cols=3)
    for i, h in enumerate(["Employee name", "Signature", "Date"]):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=10, center=True, fill=ALT_HEADER)
    for r in range(1, 6):
        for c in range(3):
            write_cell(table.rows[r].cells[c], "", size=10, fill=LIGHT_ROW if r % 2 == 0 else None)

    page_break(doc)


def rational_oven_section(doc):
    add_heading_styled(
        doc,
        "Rational Combi Oven — Equipment Risk Assessment & Operating Procedures",
        1,
    )
    try:
        add_picture(doc, "rational-oven-do-dont-care.png")
    except Exception:
        pass
    add_steps_guide(doc, "steps-rational-oven.png", "Rational oven — Step-by-step safe use")
    add_do_dont_care(
        doc,
        [
            "Stand to the side and open the door slowly to release steam",
            "Use heat-resistant gloves / dry cloths for trays and racks",
            "Select the correct programme and temperature",
            "Use only approved Rational / SK cleaning tablets as directed",
        ],
        [
            "Don’t put your face in the steam path",
            "Don’t touch hot racks or glass with bare hands",
            "Don’t overload trays or leave grease fires unattended",
            "Don’t mix oven cleaners with other chemicals",
        ],
        [
            "Check door seals and racks before service",
            "Wipe condensate and keep floors dry",
            "Run CareControl / cleaning programmes with correct tablets",
            "Keep SDS for Rational detergent and care tablets with this pack",
        ],
    )

    add_heading_styled(doc, "1. Equipment Information", 2)
    info_table(
        doc,
        [
            ("Workplace", "The Vedanta Kitchen & Retreat Centre"),
            ("Department", "Main Kitchen"),
            ("Equipment", "Rational Combi Oven (steam / convection / combi modes)"),
            ("Manufacturer / model", "Rational (model as installed on site)"),
            ("Persons completing the task", "Trained and authorised kitchen staff"),
            ("Assessor", "Shyam Prasad"),
            ("Assessment date", "27 July 2026"),
            ("Review date", "27 July 2027, or sooner after an accident, near miss, equipment change or safety concern"),
            ("Manager responsible", "[Name]"),
        ],
    )

    add_heading_styled(doc, "2. Description of Equipment", 2)
    add_body(
        doc,
        "A commercial Rational combi oven used for roasting, baking, steaming, regenerating and "
        "combination cooking. The oven can produce high temperatures, steam, hot surfaces, hot liquids "
        "and condensate. Self-cleaning / CareControl programmes may use cleaning chemicals under heat "
        "and pressure and must only be run by trained staff following manufacturer instructions.",
    )

    add_heading_styled(doc, "3. Persons at Risk", 2)
    for person in [
        "Chefs and kitchen assistants",
        "Kitchen porters and cleaning staff",
        "Maintenance contractors",
        "Staff working near the oven during service",
        "New, inexperienced or young workers",
        "Consumers, where poor cooking or cleaning controls affect food safety",
    ]:
        add_bullet(doc, person)

    add_heading_styled(doc, "4. Operating Procedure (How to Use)", 2)
    add_subheading(doc, "Before use")
    add_numbered(
        doc,
        [
            "Confirm the operator is trained and authorised on the Rational oven.",
            "Check the door seal, handle, racks, trays, probes and controls are clean and undamaged.",
            "Ensure the oven is stable, vents are clear and the surrounding floor is dry.",
            "Confirm water and power connections are correct where relevant.",
            "Select the correct mode, temperature and programme for the product.",
            "Do not operate if the door, glass, seal, racks or controls are damaged.",
        ],
    )
    add_subheading(doc, "During cooking")
    add_numbered(
        doc,
        [
            "Load trays carefully; do not overload racks or block airflow.",
            "Keep face and forearms clear when opening the door — open slowly to release steam away from the body.",
            "Announce “hot” / “steam” when opening during busy service.",
            "Use dry oven cloths or heat-resistant gloves for trays and racks.",
            "Never leave flammable materials on or against the oven.",
            "Remain aware of hot condensate and dripping liquids from trays.",
            "Stop the programme and seek help if unusual noise, smell, smoke or error codes occur.",
        ],
    )
    add_subheading(doc, "After use / cleaning")
    add_numbered(
        doc,
        [
            "Allow the oven to cool to a safe temperature before manual cleaning where required.",
            "Remove food debris from trays, racks and the cavity using approved methods.",
            "Run manufacturer cleaning / rinse programmes only with approved Rational cleaning chemicals and correct dosing.",
            "Never mix cleaning tablets/liquids with other chemicals.",
            "Wear gloves and eye protection when handling oven cleaner.",
            "Keep the door closed during automatic cleaning cycles as instructed.",
            "Wipe surrounding floors of condensate and leave the area safe for the next service.",
        ],
    )

    add_heading_styled(doc, "5. Risk Assessment", 2)
    detailed_risk_table(
        doc,
        [
            (
                "Burns from hot surfaces, trays, racks and door glass",
                "Burns to hands, arms or face",
                "15 – High",
                "Use dry heat-resistant gloves/cloths; warn others; allow cooling where practical; never touch heating elements or hot racks bare-handed.",
                "5 – Medium",
            ),
            (
                "Steam release when opening the door",
                "Scalds to face, neck, arms or chest",
                "15 – High",
                "Stand to the side; open door slowly; keep face clear of the opening; announce steam release during service.",
                "5 – Medium",
            ),
            (
                "Hot liquid / grease spill from trays",
                "Scalds, slips or secondary burns",
                "12 – High",
                "Load trays level; do not overfill; use both hands; wipe spills immediately; non-slip footwear.",
                "4 – Low",
            ),
            (
                "Entrapment / impact from door or falling racks",
                "Bruising, crush or impact injury",
                "9 – Medium",
                "Open and close door under control; seat racks fully on supports; do not overload shelves.",
                "3 – Low",
            ),
            (
                "Fire from grease build-up or unsuitable items",
                "Burns, smoke inhalation or property damage",
                "12 – High",
                "Clean regularly; never cook with excessive unattended oil; know extinguisher type and shut-off procedure; keep flammables away.",
                "4 – Low",
            ),
            (
                "Electrical fault or damaged cable / controls",
                "Electric shock, burns or fire",
                "10 – High",
                "Visual checks before use; report faults; isolate and label “DO NOT USE”; repairs by competent engineer only.",
                "5 – Medium",
            ),
            (
                "Contact with oven cleaning chemicals",
                "Skin/eye burns, irritation or breathing difficulty",
                "12 – High",
                "Use only approved Rational cleaner; follow COSHH and dosing instructions; wear gloves and eye protection; never mix chemicals; rinse as required.",
                "4 – Low",
            ),
            (
                "Slips from condensate / cleaning water around oven",
                "Slips, falls and impact injuries",
                "9 – Medium",
                "Mop condensate promptly; wet-floor signs; keep walkways clear; slip-resistant footwear.",
                "3 – Low",
            ),
            (
                "Manual handling of loaded trays / trolleys",
                "Back, shoulder or strain injuries",
                "9 – Medium",
                "Use trolley where available; keep loads manageable; ask for help with heavy pans; good posture.",
                "4 – Low",
            ),
            (
                "Undercooking / temperature abuse of food",
                "Food poisoning risk to guests",
                "12 – High",
                "Use correct programmes; probe core temperatures where required; follow HACCP and hot-holding rules.",
                "4 – Low",
            ),
            (
                "Use by untrained staff",
                "Burns, scalds, chemical injury or unsafe cooking",
                "12 – High",
                "Only trained authorised operators; supervise new staff; keep operating instructions accessible.",
                "4 – Low",
            ),
        ],
    )

    add_heading_styled(doc, "6. Required PPE", 2)
    for item in [
        "Apron / chef whites",
        "Non-slip footwear",
        "Heat-resistant gloves or dry oven cloths",
        "Chemical-resistant gloves and eye protection when handling oven cleaner",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "7. Emergency Arrangements", 2)
    for item in [
        "Burns/scalds: cool under lukewarm running water for at least 20 minutes; seek first aid; call 999 for serious burns.",
        "Fire: shut down if safe, evacuate if needed, use the correct extinguisher only if trained, call 999.",
        "Chemical contact: rinse skin/eyes as per SDS/COSHH; seek medical help for eye exposure or persistent symptoms.",
        "Faulty oven: stop programme if safe, isolate power where possible, label “DO NOT USE – FAULTY”, report to manager.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "8. Additional Notes", 2)
    for note in [
        "This section restores the Rational oven risk assessment that was missing from the earlier combined pack.",
        "Insert the exact Rational model number from the machine rating plate when known.",
        "Keep the manufacturer manual and cleaning chemical SDS with this assessment.",
    ]:
        add_bullet(doc, note)

    page_break(doc)


def chemical_risk_section(doc):
    add_heading_styled(
        doc,
        "Kitchen Chemicals (COSHH) — Sylvester Keal Product Risk Assessment",
        1,
    )
    try:
        add_picture(doc, "chemicals-coshh-do-dont-care.png")
    except Exception:
        pass
    add_steps_guide(doc, "steps-chemicals.png", "Kitchen chemicals — Step-by-step safe use")
    add_do_dont_care(
        doc,
        [
            "Wear gloves (and eye protection for concentrates / oven cleaner)",
            "Read the label and SDS before use",
            "Store chemicals locked / labelled, away from food",
            "Dilute and dose exactly as instructed",
        ],
        [
            "Never mix chemicals together",
            "Never pour chemicals into drinks or food bottles",
            "Never spray toward your face or other people",
            "Never use a product with no label or SDS",
        ],
        [
            "Keep Sylvester Keal SDS sheets with this pack",
            "Rinse food-contact surfaces after cleaning",
            "Wipe spills and wash hands before handling food",
            "Review when a new SK product is introduced",
        ],
    )

    add_heading_styled(doc, "1. Assessment Information", 2)
    info_table(
        doc,
        [
            ("Workplace", "The Vedanta Kitchen & Retreat Centre — Branston Hall, Lincoln LN4 1PD"),
            ("Company", "The Vedanta Way Limited"),
            ("Department", "Main Kitchen / Cleaning"),
            ("Activity", "Storage, dosing, use and disposal of Sylvester Keal kitchen cleaning chemicals"),
            ("Supplier", "Sylvester Keal (SK) — https://sylvesterkeal.co.uk/sk-company-brochure/"),
            ("Website", "https://thevedanta.org/"),
            ("Persons completing the task", "Chefs, kitchen porters, cleaning staff and authorised users"),
            ("Assessor", "Shyam Prasad"),
            ("Assessment date", "27 July 2026"),
            ("Review date", "27 July 2027, or sooner after an incident, product change or SDS update"),
            ("Manager responsible", "[Name]"),
        ],
    )

    add_heading_styled(doc, "2. Site Chemical Inventory (Sylvester Keal brochure)", 2)
    add_body(
        doc,
        "Products taken from the official Sylvester Keal Product Guide (Machine & Manual Dishwashing Products "
        "and Oven Cleaning Products pages). Original brochure product photography is included below. "
        "All SK products have COSHH Safety Data Sheets — keep the current SDS for each code with this pack. "
        "Brochure: https://sylvesterkeal.co.uk/sk-company-brochure/",
    )

    add_subheading(doc, "Original brochure — Machine & Manual Dishwashing Products")
    try:
        add_picture(doc, "sk-dishwashing-products.png", width_inches=6.3)
    except Exception:
        try:
            add_picture(doc, "sk-dishwashing-brochure-spread.png", width_inches=6.3)
        except Exception:
            pass

    products = [
        ("SK Premium Dishwashing Detergent", "1A", "2×5L / 20L", "Hard-water machine detergent; tannin & scale control"),
        ("SK Premium Dishwashing Rinse Aid", "2B", "2×5L / 20L", "Machine rinse aid for hygienic streak-free finish"),
        ("SK Machine Destainer", "2HX", "2×5L / 3×5L / 20L", "Tea/coffee destainer with sanitiser for machines"),
        ("SK Bacti Manual Dishwashing Detergent", "3AB", "2×5L / 20L", "Manual dishwashing; cuts grease; streak-free rinse"),
        ("SK Super Lemon Sinkwash", "3L", "2×5L / 12×1L", "Manual wash-up & light duty general cleaning"),
        ("SK Super Lemon Sinkwash (1L)", "3L1", "12×1L", "1L sinkwash bottle for manual washing"),
        ("SK Salt Granular (Hydrosoft)", "12705", "25KG", "Granular vacuum salt for water softeners"),
        ("SK Salt Pebble", "PE12705", "25KG", "Salt tablets for water softeners / limescale prevention"),
        ("SK 1DT Dishwasher Tablets", "1DT", "15×7 tablets", "Clean n Fresh 7-in-1 lemon dishwasher tablets"),
    ]
    table = doc.add_table(rows=1 + len(products), cols=4)
    for i, h in enumerate(["Product (brochure)", "Code", "Pack size", "Intended use"]):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=8, center=True, fill=HEADER_BG)
    for r, row in enumerate(products, 1):
        fill = LIGHT_ROW if r % 2 == 0 else None
        for c, val in enumerate(row):
            write_cell(table.rows[r].cells[c], val, size=8, fill=fill)
    doc.add_paragraph()

    add_subheading(doc, "Original brochure — Oven Cleaning Products")
    try:
        add_picture(doc, "sk-oven-cleaner-products.png", width_inches=6.3)
    except Exception:
        try:
            add_picture(doc, "sk-oven-cleaner-brochure-spread.png", width_inches=6.3)
        except Exception:
            pass

    oven_products = [
        ("SK Hot Oven Cleaner", "7A750", "6×750ml", "Trigger spray for ovens, hot plates, grills, griddles"),
        ("SK Hot Oven Cleaner", "7A", "2×5L", "5L hot oven cleaner for commercial cooking equipment"),
        ("SK Carbon Remover", "7C", "10KG", "Caustic powder for burnt-on oils/fats on trays & equipment"),
        ("SK Combi Oven Rinse Aid", "2K", "10L", "Integral combi-oven rinse system (auto dosing)"),
        ("SK Combi Oven Cleaner", "7K", "10L", "Integral combi-oven cleaner / degreaser (auto dosing)"),
        ("SK Rational Detergent Tablets", "OCA8294", "1×100", "Red Rational combi oven detergent tablets"),
        ("SK Rational Active Green Care Cleaner Tablets", "56.01.535", "1×150", "iCombi Pro / Classic phosphate-free care cleaner"),
        ("SK Rational Combi Care Control Tablets", "OCA8357", "1×150", "Care/control tablets for Rational & Lincat SCC"),
    ]
    table = doc.add_table(rows=1 + len(oven_products), cols=4)
    for i, h in enumerate(["Product (brochure)", "Code", "Pack size", "Intended use"]):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=8, center=True, fill=HEADER_BG)
    for r, row in enumerate(oven_products, 1):
        fill = LIGHT_ROW if r % 2 == 0 else None
        for c, val in enumerate(row):
            write_cell(table.rows[r].cells[c], val, size=8, fill=fill)
    doc.add_paragraph()

    add_body(
        doc,
        "Source images: Sylvester Keal Company Brochure 2025 — Machine & Manual Dishwashing Products "
        "and Oven Cleaning Products pages (original product photography).",
        size=9,
    )

    add_heading_styled(doc, "3. Step-by-step — Dishwashing chemicals", 2)
    add_subheading(doc, "Machine dishwashing (1A / 2B / salt / tablets)")
    add_numbered(
        doc,
        [
            "Check dishwasher is empty of food debris and spray arms are free.",
            "Ensure softener salt (pebble or granular) is topped up if the machine requires it.",
            "Confirm detergent (1A or tablets 1DT) and rinse aid (2B) are correctly loaded / dosed.",
            "Wear gloves when handling concentrates or tablets.",
            "Run the cycle; do not open mid-cycle if steam/chemical vapour is present.",
            "On empty containers, store upright, labelled, away from food.",
        ],
    )
    add_subheading(doc, "Manual dishwashing (Super Lemon Sinkwash / Bacti Manual)")
    add_numbered(
        doc,
        [
            "Scrape food soil into waste before washing.",
            "Dilute sinkwash / Bacti Manual as per label — do not guess a stronger mix.",
            "Wash, rinse and sanitise following kitchen procedure.",
            "Change water when heavily soiled; wipe spills around the sink.",
            "Wash hands after chemical contact and before handling clean service ware.",
        ],
    )
    add_subheading(doc, "Oven / heavy degreasing (Hot Oven Cleaner / Carbon Remover / Rational tablets / Combi 7K & 2K)")
    add_numbered(
        doc,
        [
            "Only trained staff may use Hot Oven Cleaner (7A/7A750), Carbon Remover (7C), Combi Oven Cleaner/Rinse (7K/2K) or Rational tablets.",
            "Wear chemical gloves and eye protection.",
            "For Rational: use only the correct tablet type for the oven model (OCA8294 / 56.01.535 / OCA8357); run the manufacturer cleaning programme; keep door closed.",
            "For Hot Oven Cleaner: follow label; do not use on aluminium or zinc alloys where the product warns against it.",
            "Never mix these products with sinkwash, bleach or other cleaners.",
            "Rinse food-contact surfaces thoroughly after use.",
        ],
    )

    add_heading_styled(doc, "4. Persons at Risk", 2)
    for person in [
        "Kitchen porters and cleaning staff",
        "Chefs and assistants using chemicals during service or close-down",
        "Maintenance staff",
        "Other employees nearby during spraying or mopping",
        "Guests, if residues remain on food-contact surfaces",
        "New or inexperienced workers",
    ]:
        add_bullet(doc, person)

    add_heading_styled(doc, "5. Risk Assessment", 2)
    detailed_risk_table(
        doc,
        [
            (
                "Skin contact with concentrates (1A, sinkwash, degreaser, oven cleaner)",
                "Irritation, dermatitis or chemical burns",
                "12 – High",
                "Gloves for concentrates; correct dilution; wash splashes promptly; report skin problems early.",
                "4 – Low",
            ),
            (
                "Eye contact from splash (oven cleaner / caustic degreaser / dosing)",
                "Eye irritation or serious eye damage",
                "12 – High",
                "Eye protection for aggressive cleaners; pour carefully; know eye-wash location.",
                "4 – Low",
            ),
            (
                "Mixing incompatible SK / other chemicals",
                "Toxic gas, burns or violent reaction",
                "15 – High",
                "Never mix products; rinse between different cleaners; train and supervise staff.",
                "4 – Low",
            ),
            (
                "Rational tablet / hot oven cleaner misuse",
                "Severe skin/eye burns or inhalation injury",
                "15 – High",
                "Trained staff only; correct tablet for model; door closed on auto clean; PPE; follow SDS.",
                "5 – Medium",
            ),
            (
                "Chemical in unmarked / food bottle",
                "Accidental ingestion or food contamination",
                "15 – High",
                "Original SK labelled containers only; separate chemical store; stock checks.",
                "4 – Low",
            ),
            (
                "Residue on crockery or prep surfaces",
                "Chemical contamination of food",
                "12 – High",
                "Correct rinse aid dosing; rinse food-contact surfaces; follow contact times for sanitisers.",
                "4 – Low",
            ),
            (
                "Spillage of detergent / degreaser",
                "Slips and skin/eye exposure",
                "9 – Medium",
                "Contain with PPE; wet-floor sign; dispose correctly; report significant spills.",
                "3 – Low",
            ),
            (
                "Untrained person using chemicals",
                "Incorrect use leading to injury or contamination",
                "12 – High",
                "COSHH induction; product-specific instruction; SDS accessible; supervise until competent.",
                "4 – Low",
            ),
        ],
    )

    add_heading_styled(doc, "6. Required PPE", 2)
    for item in [
        "Chemical-resistant gloves suitable for the product",
        "Eye protection for concentrates, Hot Oven Cleaner, Heavy Duty Degreaser and Rational cleaning chemicals",
        "Apron where specified",
        "Non-slip footwear",
        "Additional PPE as stated on the Sylvester Keal product SDS",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "7. Emergency Arrangements", 2)
    for item in [
        "Skin contact: remove contaminated clothing; rinse with plenty of water; seek first aid / medical advice if irritation persists.",
        "Eye contact: rinse immediately with clean water for at least 15 minutes; seek urgent medical attention.",
        "Inhalation: move to fresh air; seek medical help if breathing is affected.",
        "Swallowing: do not induce vomiting unless SDS advises; call 999 / poison guidance and take the product label/SDS.",
        "Large spill or unknown reaction: evacuate immediate area, ventilate if safe, inform manager, do not neutralise by mixing other chemicals.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "8. Training & Records", 2)
    for item in [
        "All users must receive COSHH awareness and product-specific instruction before unsupervised use.",
        "Keep Sylvester Keal SDS sheets for every chemical code in an accessible folder.",
        "Record training, product changes and any chemical-related incidents or near misses.",
        "Review this assessment when a new SK product is introduced or an SDS is updated.",
        "Brochure reference: https://sylvesterkeal.co.uk/sk-company-brochure/",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "9. Additional Notes", 2)
    for note in [
        "This inventory is based on products confirmed by kitchen knowledge and the Sylvester Keal brochure range.",
        "Where a product SDS sets stricter controls than this sheet, follow the SDS.",
        "Rational cleaning tablets must only be used in compatible Rational / approved ovens.",
    ]:
        add_bullet(doc, note)

    page_break(doc)


def sign_off(doc):
    add_heading_styled(doc, "Document Control & Sign-Off", 1)
    add_body(
        doc,
        "This combined pack replaces the separate equipment and area risk assessment files listed on the cover. "
        "It should be reviewed after any incident, equipment change, or at least annually.",
    )
    info_table(
        doc,
        [
            ("Document owner", "Shyam Prasad"),
            ("Approved by", ""),
            ("Approval date", ""),
            ("Next review", "27/07/2027"),
            ("Version", "2.3 Added fryer, dough mixer, blast chiller, bratt pan, Wrapmaster, induction & tray storage"),
        ],
    )
    add_heading_styled(doc, "Staff acknowledgement", 2)
    add_body(doc, "I confirm I have read and understood the relevant sections for equipment and tasks I use.")

    table = doc.add_table(rows=6, cols=3)
    for i, h in enumerate(["Name", "Role / Area", "Date / Signature"]):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=10, center=True, fill=ALT_HEADER)
    for r in range(1, 6):
        for c in range(3):
            write_cell(table.rows[r].cells[c], "", size=10, fill=LIGHT_ROW if r % 2 == 0 else None)


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    from extra_equipment import EXTRA_SECTION_TITLES, add_all_extra_equipment

    section_titles = [
        "Electrolux / Dito Sama TRK70 Combined Cutter & Vegetable Slicer",
        "Rational Combi Oven",
        "Kitchen Chemicals (COSHH) — Sylvester Keal",
        "Thermomix TM6",
        "Caso Ice Creamer",
        "Ninja Hand Blender",
        "Mixers & Blenders — Visual Safety Guide",
        "KitchenAid Professional Stand Mixer",
        "Waring Stick Blender",
        "Knives & Mandoline — Visual Safety Guide",
        "Knives, Mandoline, Slicers & Peelers",
        *EXTRA_SECTION_TITLES,
        "Vedanta Kitchen Risk Assessment",
        "Vedanta Catering Services Risk Assessment",
        "Vedanta Front of House & Kitchen Risk Assessment",
        "Vedanta Generic Kitchen Risk Assessment (June 2024)",
        "Document Control & Sign-Off",
    ]

    cover_page(doc)
    contents_page(doc, section_titles)

    # --- Electrolux / Dito Sama TRK70 ---
    trk70_section(doc)

    # --- Rational oven ---
    rational_oven_section(doc)

    # --- Kitchen chemicals / COSHH ---
    chemical_risk_section(doc)

    # --- Thermomix TM6 ---
    equipment_section(
        doc,
        "Thermomix TM6 — Equipment Risk Assessment & Operating Procedures",
        "Thermomix",
        "TM6",
        "A multi-function cooking appliance used for chopping, blending, cooking, steaming, "
        "kneading and precise heating. Used for soups, sauces, doughs, purees and prepared dishes "
        "in bakery and catering production.",
        [
            "Place the Thermomix on a stable, dry, level work surface with adequate ventilation around the base.",
            "Ensure the mixing bowl is correctly seated and locked before use.",
            "Fit the correct blade or accessory for the task and secure the lid and measuring cup as required.",
            "Select the appropriate mode, temperature and speed for the recipe; do not exceed recommended fill levels.",
            "Keep hands, utensils and cloths away from the blade area while the machine is running.",
            "Allow hot contents to settle before opening the lid; open away from the face to avoid steam.",
            "Use the spatula only through the lid opening when the manufacturer instructions allow.",
            "Switch off and unplug before changing accessories, emptying the bowl or cleaning.",
            "Clean food-contact parts thoroughly after each use and follow HACCP hygiene controls.",
        ],
        [
            (
                "Blade contact / finger injury",
                "Staff",
                "High",
                "Never reach into the bowl while powered; lock lid; unplug before removing blade.",
                "Low",
            ),
            (
                "Steam / hot liquid burns",
                "Staff",
                "High",
                "Use correct temperature settings; open lid carefully; use oven cloths for hot bowls.",
                "Low",
            ),
            (
                "Electrical shock",
                "Staff",
                "Medium",
                "Inspect cable and base; keep area dry; unplug before cleaning.",
                "Low",
            ),
            (
                "Splash / ejection of hot contents",
                "Staff",
                "Medium",
                "Do not overfill; use measuring cup; start at low speed for hot liquids.",
                "Low",
            ),
            (
                "Food contamination",
                "Guests / Staff",
                "High",
                "Clean and sanitise bowl, lid and blade; separate raw and ready-to-eat use; follow HACCP.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Heat-resistant gloves or cloths when handling hot bowls / steam",
        ],
        steps_image="steps-thermomix.png",
    )

    # --- Caso Ice Creamer ---
    equipment_section(
        doc,
        "Caso Ice Creamer — Equipment Risk Assessment & Operating Procedures",
        "Caso",
        "Ice Creamer",
        "A countertop ice cream / frozen dessert machine used to churn and freeze ice cream, "
        "sorbet and similar products for catering and retreat service.",
        [
            "Position the machine on a stable, ventilated surface away from heat sources.",
            "Ensure the freezing bowl / canister is correctly pre-frozen or pre-chilled as per the manual.",
            "Assemble the paddle, lid and bowl securely before switching on.",
            "Prepare mix to the correct recipe and temperature; do not overfill the bowl.",
            "Start the machine and supervise during churning; do not leave unattended for long periods.",
            "Do not insert hands or utensils into the bowl while the paddle is moving.",
            "Switch off before removing the paddle or scraping the finished product.",
            "Transfer product with clean utensils into sanitised containers and store under correct temperature control.",
            "Clean all food-contact parts after use and dry thoroughly before storage.",
        ],
        [
            (
                "Finger entrapment in paddle",
                "Staff",
                "High",
                "Keep hands clear while running; switch off before scraping or removing paddle.",
                "Low",
            ),
            (
                "Cold burns / freezer bowl handling",
                "Staff",
                "Medium",
                "Handle frozen canister with dry cloths or gloves; do not touch with wet hands.",
                "Low",
            ),
            (
                "Electrical shock",
                "Staff",
                "Medium",
                "Keep cable and controls dry; inspect before use; unplug before cleaning.",
                "Low",
            ),
            (
                "Slips from spilled mix",
                "Staff",
                "Medium",
                "Wipe spills immediately; keep floor clear around machine.",
                "Low",
            ),
            (
                "Food contamination / temperature abuse",
                "Guests / Staff",
                "High",
                "Use clean utensils; chill finished product promptly; follow shelf-life and HACCP controls.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Protective gloves when handling frozen canisters",
        ],
        steps_image="steps-caso-icecreamer.png",
    )

    # --- Ninja Hand Blender ---
    equipment_section(
        doc,
        "Ninja Hand Blender — Equipment Risk Assessment & Operating Procedures",
        "Ninja",
        "Hand Blender",
        "A handheld immersion blender used for pureeing soups, sauces, dressings and soft mixtures "
        "directly in pans or containers.",
        [
            "Check the blender shaft, blade guard and cable before use; do not use if damaged.",
            "Ensure the pan or jug is stable and only partially filled to avoid splash-over.",
            "Fully immerse the blade guard in the food before switching on.",
            "Start on a low speed and keep the head submerged while blending.",
            "Never blend with the blade above the liquid surface or near the rim of a shallow pan.",
            "Do not insert fingers or utensils near the blade while powered.",
            "Switch off and unplug before removing the shaft or cleaning the blade.",
            "Allow hot liquids to cool slightly where practical to reduce splash risk.",
            "Wash detachable parts, sanitise, dry and store with blade guard protected.",
        ],
        [
            (
                "Blade cuts",
                "Staff",
                "High",
                "Keep hands clear; unplug before cleaning; store with guard / cover.",
                "Low",
            ),
            (
                "Hot liquid splash / burns",
                "Staff",
                "High",
                "Immerse fully before starting; blend gently; use tall containers where possible.",
                "Low",
            ),
            (
                "Electrical shock",
                "Staff",
                "Medium",
                "Keep motor body dry; never immerse the handle; inspect cable.",
                "Low",
            ),
            (
                "Dropped appliance / impact injury",
                "Staff",
                "Medium",
                "Use two hands for deep pans; dry grip; do not stretch cables across walkways.",
                "Low",
            ),
            (
                "Food contamination",
                "Guests / Staff",
                "High",
                "Clean and sanitise shaft and blade after each use; avoid cross-contamination.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Cut-resistant glove recommended when washing the blade",
        ],
        steps_image="steps-ninja-blender.png",
    )

    # --- Mixers & blenders visual guide ---
    add_heading_styled(doc, "Mixers & Blenders — Visual Safety Guide", 1)
    try:
        add_picture(doc, "mixers-blenders-do-dont-care.png")
    except Exception:
        pass
    add_do_dont_care(
        doc,
        [
            "Start stand mixers on the lowest speed",
            "Immerse stick-blender blade fully before switching on",
            "Unplug before changing tools or washing blades",
            "Use stable bowls / pans and dry grip",
        ],
        [
            "Don’t put spatulas or hands in a moving bowl",
            "Don’t leave a plugged-in stick blender unattended in a pot",
            "Don’t wash blades while still connected to power",
            "Don’t overfill bowls or shallow pans",
        ],
        [
            "Inspect cables and attachments before use",
            "Clean and sanitise food-contact parts after each product",
            "Store blades with guards / covers",
            "Remove damaged equipment from service",
        ],
    )
    page_break(doc)

    # --- KitchenAid (from user's screenshot content) ---
    equipment_section(
        doc,
        "KitchenAid Professional Stand Mixer — Equipment Risk Assessment & Operating Procedures",
        "KitchenAid",
        "Professional Stand Mixer",
        "A heavy-duty stand mixer used for mixing doughs, batters, creams and other food ingredients. "
        "Commonly used for bakery and catering production.",
        [
            "Ensure the mixer is placed on a stable, dry surface before use.",
            "Attach the correct tool (dough hook, whisk or flat beater) and lock the bowl securely.",
            "Add ingredients carefully; do not exceed the recommended bowl capacity.",
            "Start the mixer on the lowest speed and increase gradually as needed.",
            "Never insert hands, spatulas or utensils into the bowl while the mixer is operating.",
            "Stop the mixer completely before scraping the bowl or changing tools.",
            "Unplug the mixer before cleaning or removing attachments.",
            "Clean all food-contact parts thoroughly after each use.",
            "Store attachments safely when not in use.",
        ],
        [
            (
                "Finger entrapment",
                "Staff",
                "High",
                "Keep hands away from moving parts; ensure mixer is off when changing tools.",
                "Low",
            ),
            (
                "Electrical shock",
                "Staff",
                "Medium",
                "Inspect cables; keep hands dry; unplug before cleaning.",
                "Low",
            ),
            (
                "Cuts from whisk or hook edges",
                "Staff",
                "Medium",
                "Handle attachments carefully; unplug before removal.",
                "Low",
            ),
            (
                "Mixer movement during heavy loads",
                "Staff",
                "Medium",
                "Ensure stable surface; supervise during thick dough mixing.",
                "Low",
            ),
            (
                "Food contamination",
                "Guests / Staff",
                "High",
                "Clean thoroughly; sanitise food-contact parts; follow HACCP.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Heat-resistant gloves (when working with warm mixtures)",
        ],
        steps_image="steps-kitchenaid.png",
    )

    # --- Waring Stick Blender ---
    equipment_section(
        doc,
        "Waring Stick Blender — Equipment Risk Assessment & Operating Procedures",
        "Waring",
        "Stick / Immersion Blender",
        "A commercial stick blender used for blending soups, sauces and large-batch liquids "
        "in pots and deep containers within the kitchen.",
        [
            "Inspect the shaft, blade, switch and power cable before each use.",
            "Place the pot on a stable surface; do not overfill.",
            "Immerse the blade fully before switching on.",
            "Hold the blender with a firm grip; support the pot if needed.",
            "Blend at controlled speed to avoid splash-back of hot liquids.",
            "Never place hands near the blade while the unit is plugged in.",
            "Switch off and unplug before dismantling or washing.",
            "Clean immediately after use; sanitise food-contact surfaces.",
            "Store hanging or in a designated safe location with blade protected.",
        ],
        [
            (
                "Severe blade cuts",
                "Staff",
                "High",
                "Unplug before cleaning; never leave plugged-in blender unattended in a pot; use blade cover in storage.",
                "Low",
            ),
            (
                "Hot liquid burns / splash",
                "Staff",
                "High",
                "Start submerged; use deep pots; reduce speed for hot liquids.",
                "Low",
            ),
            (
                "Electrical shock",
                "Staff",
                "Medium",
                "Keep motor housing dry; check cable integrity; RCD-protected circuits where available.",
                "Low",
            ),
            (
                "Musculoskeletal strain",
                "Staff",
                "Medium",
                "Use correct posture; avoid extended one-handed use; take breaks on large batches.",
                "Low",
            ),
            (
                "Food contamination",
                "Guests / Staff",
                "High",
                "Clean and sanitise after each product; follow allergen and HACCP controls.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Cut-resistant glove when washing blades",
        ],
        extra_notes=[
            "This section consolidates Risk_Assessment_Waring_Stick_Blender.docx and the duplicate “(1)” file.",
        ],
        steps_image="steps-waring-blender.png",
    )

    # --- Knives visual guide ---
    add_heading_styled(doc, "Knives & Mandoline — Visual Safety Guide", 1)
    try:
        add_picture(doc, "knives-do-dont-care.png")
    except Exception:
        pass
    add_steps_guide(doc, "steps-knives-peelers.png", "Knives & peelers — Step-by-step safe use")
    add_do_dont_care(
        doc,
        [
            "Use claw grip and a stable board with non-slip mat",
            "Use hand guard or cut-resistant glove on mandoline / slicer",
            "Keep knives and peelers in the designated kitchen drawer only",
            "Wash knives and peelers individually and carefully",
        ],
        [
            "Don’t catch a falling knife",
            "Don’t push food on a mandoline with bare fingers",
            "Don’t leave knives or peelers hidden in sink water",
            "Don’t rummage blindly in the knife/peeler drawer",
        ],
        [
            "Store knives and peelers neatly in the drawer with blades protected / facing one way",
            "Open the drawer carefully and pick up tools by the handle",
            "Use colour-coded boards for raw / ready-to-eat / allergens",
            "Report damaged tools and remove from service",
        ],
    )
    page_break(doc)

    # --- Knives / Mandoline / Slicers / Peelers ---
    equipment_section(
        doc,
        "Knives, Mandoline, Slicers & Peelers — Risk Assessment & Safe Use",
        "Various",
        "Chef knives, mandoline, slicers and peelers",
        "Hand tools and small cutting equipment used for preparation of fruit, vegetables, meat "
        "and other ingredients. At The Vedanta Kitchen, kitchen knives and peelers are kept in the "
        "designated drawer. This section merges the Knives/Mandoline/Peelers and "
        "Knives/Slicers/Peelers assessments into one procedure.",
        [
            "Select the correct knife or peeler for the task; check blades are clean and undamaged.",
            "Use a stable chopping board with a non-slip mat underneath.",
            "Keep fingers curled and clear of the blade path (claw grip).",
            "For mandoline / slicers: always use the hand guard or cut-resistant glove; never push product by hand.",
            "Cut away from the body; do not catch falling knives — let them drop.",
            "Wash knives and peelers individually; never leave blades loose in sinks under water.",
            "Return knives and peelers to the designated kitchen drawer after cleaning and drying.",
            "In the drawer: store neatly with handles accessible, blades protected or facing one direction — do not leave blades exposed and jumbled.",
            "Open the drawer carefully; pick tools up by the handle only — never rummage with fingers among blades.",
            "Sharpen / hone as required; a dull blade increases slip risk.",
            "Report damaged tools and remove from service.",
        ],
        [
            (
                "Cuts / lacerations from knives",
                "Staff",
                "High",
                "Training on safe knife skills; claw grip; cut-resistant gloves for high-risk tasks; careful drawer handling.",
                "Low",
            ),
            (
                "Severe cuts from mandoline / slicer",
                "Staff",
                "High",
                "Mandatory use of guard or cut-resistant glove; dismantle carefully for cleaning.",
                "Low",
            ),
            (
                "Cuts when taking knives or peelers from the drawer",
                "Staff",
                "High",
                "Designated knife/peeler drawer only; store neatly with handles out / blades protected; open carefully; pick up by handle; no blind rummaging.",
                "Low",
            ),
            (
                "Cross-contamination (raw / allergens)",
                "Guests / Staff",
                "High",
                "Colour-coded boards and knives; wash between products; follow HACCP and allergen matrix.",
                "Low",
            ),
            (
                "Slips from peelings / waste on floor",
                "Staff",
                "Medium",
                "Clear waste regularly; wipe floors; use non-slip footwear.",
                "Low",
            ),
            (
                "Incorrect tool storage injuries",
                "Staff",
                "Medium",
                "Keep knives and peelers only in the designated drawer; never leave in sinks or on open worktops after use.",
                "Low",
            ),
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Cut-resistant gloves for mandoline, slicer and high-volume prep",
        ],
        extra_notes=[
            "Site practice: kitchen knives and peelers are stored in the designated drawer.",
            "If blade guards/sheaths are available, fit them before returning knives to the drawer.",
        ],
    )

    # --- Site equipment from kitchen photos ---
    add_all_extra_equipment(doc)

    # --- Kitchen area ---
    area_section(
        doc,
        "Vedanta Kitchen Risk Assessment",
        "Covers general kitchen operations including food preparation, cooking, cleaning, "
        "storage, waste handling and staff movement within The Vedanta Kitchen.",
        [
            (
                "Slips, trips and falls",
                "Staff / Visitors",
                "High",
                "Clean spills immediately; clear walkways; non-slip mats; suitable footwear.",
                "Low",
            ),
            (
                "Burns from ovens, hobs, hot pans",
                "Staff",
                "High",
                "Use dry cloths / gloves; announce “hot”; keep pan handles turned inward.",
                "Low",
            ),
            (
                "Fire (oil, electrical, gas)",
                "Staff / Guests",
                "High",
                "Never leave fryers/pans unattended; know extinguisher types; fire training; extract cleaning.",
                "Low",
            ),
            (
                "Manual handling injuries",
                "Staff",
                "Medium",
                "Team lifts for heavy stock; use trolleys; training on safe lifting.",
                "Low",
            ),
            (
                "Food poisoning / contamination",
                "Guests",
                "High",
                "Temperature control; date labelling; cleaning schedules; allergen controls; HACCP.",
                "Low",
            ),
            (
                "Chemical exposure (cleaners)",
                "Staff",
                "Medium",
                "COSHH sheets; correct dilution; PPE; store chemicals away from food.",
                "Low",
            ),
        ],
        [
            "Maintain daily opening/closing cleaning checklists.",
            "Record fridge/freezer temperatures at required intervals.",
            "Keep first-aid kit and burn gel accessible.",
            "Report all near-misses and incidents.",
        ],
        [
            "Apron / chef whites",
            "Non-slip footwear",
            "Heat-resistant gloves / cloths as needed",
            "Gloves for cleaning chemicals",
        ],
    )

    # --- Catering services ---
    area_section(
        doc,
        "Vedanta Catering Services Risk Assessment",
        "Covers off-site and on-site catering service: transport of food, service set-up, "
        "hot holding, buffet/service lines, and breakdown.",
        [
            (
                "Temperature abuse during transport",
                "Guests",
                "High",
                "Use insulated boxes; hot-hold above 63°C or cold below 5°C; temperature logs.",
                "Low",
            ),
            (
                "Vehicle loading injuries / traffic risk",
                "Staff",
                "Medium",
                "Use trolleys; secure loads; park safely; two-person lifts for heavy items.",
                "Low",
            ),
            (
                "Burns / hot holding equipment",
                "Staff",
                "Medium",
                "Warn guests/staff of hot surfaces; use correct lids and gloves.",
                "Low",
            ),
            (
                "Allergen cross-contact at service",
                "Guests",
                "High",
                "Clear allergen labelling; separate utensils; trained service staff; guest enquiry process.",
                "Low",
            ),
            (
                "Slips at service points",
                "Staff / Guests",
                "Medium",
                "Manage cable runs; wipe spills; clear guest walkways.",
                "Low",
            ),
        ],
        [
            "Complete a service checklist before departure and on arrival.",
            "Keep probe thermometer and sanitising wipes with the catering kit.",
            "Return leftovers under temperature control or dispose per food safety policy.",
        ],
        [
            "Apron / service uniform",
            "Non-slip footwear",
            "Heat-resistant gloves for hot equipment",
            "Disposable gloves for ready-to-eat service where required",
        ],
    )

    # --- FOH and Kitchen ---
    area_section(
        doc,
        "Vedanta Front of House & Kitchen Risk Assessment",
        "Covers interaction between front-of-house and kitchen areas: service pass, guest areas, "
        "beverage service, and shared walkways during service periods.",
        [
            (
                "Collision in pass / corridor",
                "Staff",
                "Medium",
                "Keep-left / call-outs (“behind”, “hot”); keep floors clear; limit congestion.",
                "Low",
            ),
            (
                "Hot plate / spill burns at pass",
                "Staff",
                "Medium",
                "Use dry cloths; communicate plate status; stable stacking.",
                "Low",
            ),
            (
                "Guest slips in dining areas",
                "Guests / Staff",
                "Medium",
                "Prompt spill response; suitable mats at entrances; clear trailing cables.",
                "Low",
            ),
            (
                "Broken glass / crockery",
                "Staff / Guests",
                "Medium",
                "Clear immediately with dustpan; dispose in glass bin; cordon if needed.",
                "Low",
            ),
            (
                "Allergen information failure",
                "Guests",
                "High",
                "Up-to-date allergen matrix; FOH trained to check with kitchen; never guess.",
                "Low",
            ),
        ],
        [
            "FOH and kitchen brief before busy service.",
            "Keep emergency exits and fire routes clear at all times.",
            "Ensure allergen folder / matrix is current for the day’s menu.",
        ],
        [
            "Non-slip footwear",
            "Service apron where required",
            "Gloves for clearing waste / glass",
        ],
    )

    # --- Generic kitchen June 2024 ---
    area_section(
        doc,
        "Vedanta Generic Kitchen Risk Assessment (June 2024)",
        "General kitchen hazards and standard controls for The Vedanta Kitchen & Retreat Centre, "
        "based on the June 2024 generic kitchen risk assessment. Use alongside equipment-specific sheets.",
        [
            (
                "General kitchen injury (cuts, burns, strains)",
                "Staff",
                "High",
                "Induction training; supervision of new staff; safe systems of work; PPE.",
                "Low",
            ),
            (
                "Electrical equipment faults",
                "Staff",
                "Medium",
                "Visual checks before use; PAT testing schedule; remove defective kit from service.",
                "Low",
            ),
            (
                "Poor housekeeping",
                "Staff / Guests",
                "Medium",
                "Clean-as-you-go; end-of-day deep clean; designated storage for tools and stock.",
                "Low",
            ),
            (
                "Pest activity",
                "Guests / Staff",
                "Medium",
                "Seal food; waste control; report sightings; pest contractor visits.",
                "Low",
            ),
            (
                "Inadequate first aid response",
                "Staff / Guests",
                "Medium",
                "Trained first aiders; stocked kit; incident reporting forms available.",
                "Low",
            ),
        ],
        [
            "Review this generic assessment whenever menus, staffing or layout change.",
            "Keep SDS / COSHH information accessible for all cleaning chemicals.",
            "Link this pack to staff induction and refresher training records.",
        ],
        [
            "Apron",
            "Non-slip footwear",
            "Task-specific gloves (heat, cut, chemical)",
        ],
        notes=[
            "Source file reference: Vedanta generic kitchen risk assessments June 2024 (AutoRecovered).",
            "Where an equipment sheet and this generic sheet both apply, follow the stricter control.",
        ],
    )

    sign_off(doc)

    out = "/workspace/vedanta-risk-assessments/Vedanta_Combined_Kitchen_Risk_Assessments.docx"
    brochure = "/workspace/vedanta-risk-assessments/Vedanta_Kitchen_Safety_Brochure_Risk_Pack.docx"
    doc.save(out)
    doc.save(brochure)
    print(out)
    print(brochure)


if __name__ == "__main__":
    build()
