#!/usr/bin/env python3
"""Separate staff kitchen-access risk assessment for The Vedanta Way Limited.

This is intentionally separate from the equipment risk pack so each department
can receive a focused copy for anyone entering the kitchen.
"""

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

IMG_DIR = "/workspace/vedanta-risk-assessments/images"
BRAND_DIR = f"{IMG_DIR}/brand"
STAFF_DIR = f"{IMG_DIR}/staff"

NAVY = RGBColor(0x1A, 0x1A, 0x1A)
GOLD = RGBColor(0x8A, 0x73, 0x4A)
DARK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_ROW = "F5F2EC"
HEADER_BG = "1A1A1A"
ALT_HEADER = "8A734A"


def set_run(run, *, size=11, bold=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def write_cell(cell, text, *, bold=False, color=DARK, size=10, center=False, fill=None):
    clear_paragraph(cell.paragraphs[0])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, bold=True, color=NAVY, font="Georgia")
    else:
        set_run(run, size=12, bold=True, color=GOLD, font="Georgia")
    return p


def add_body(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    clear_paragraph(p)
    run = p.add_run(text)
    set_run(run, size=11, color=DARK)
    p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(f"{i}. {item}")
        set_run(run, size=11, color=DARK)


def info_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        write_cell(table.rows[i].cells[0], label, bold=True, size=10, fill=LIGHT_ROW, color=NAVY)
        write_cell(table.rows[i].cells[1], value, size=10)
    doc.add_paragraph()


def risk_table(doc, hazards):
    table = doc.add_table(rows=1 + len(hazards), cols=5)
    headers = ["Hazard", "Who is at risk", "Risk before", "Control measures", "Risk after"]
    for i, h in enumerate(headers):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=8, center=True, fill=HEADER_BG)
    for r, row in enumerate(hazards, 1):
        fill = LIGHT_ROW if r % 2 == 0 else None
        for c, val in enumerate(row):
            write_cell(
                table.rows[r].cells[c],
                val,
                size=8,
                center=(c in (2, 4)),
                fill=fill,
                bold=(c in (2, 4)),
                color=GOLD if c in (2, 4) else DARK,
            )
    doc.add_paragraph()


def add_picture(doc, filename, width_inches=6.3):
    path = None
    for candidate in (f"{STAFF_DIR}/{filename}", f"{BRAND_DIR}/{filename}", f"{IMG_DIR}/{filename}"):
        if os.path.exists(candidate):
            path = candidate
            break
    if not path:
        add_body(doc, f"(Image unavailable: {filename})", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(path, width=Inches(width_inches))


def page_break(doc):
    doc.add_page_break()


def add_do_dont_care(doc, do_items, dont_items, care_items):
    add_heading(doc, "Do / Don’t / Care", 2)
    add_body(doc, "DO", bold=True, space_after=2)
    for item in do_items:
        add_bullet(doc, item)
    add_body(doc, "DON’T", bold=True, space_after=2)
    for item in dont_items:
        add_bullet(doc, item)
    add_body(doc, "CARE", bold=True, space_after=2)
    for item in care_items:
        add_bullet(doc, item)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Cover
    try:
        add_picture(doc, "Vedanta-Way-Ltd-text-banner.jpg", 6.5)
    except Exception:
        pass
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("THE VEDANTA WAY LIMITED")
    set_run(run, size=12, bold=True, color=GOLD, font="Georgia")

    main = doc.add_paragraph()
    main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main.add_run("Staff Kitchen Access\nRisk Assessment & Rules")
    set_run(run, size=24, bold=True, color=NAVY, font="Georgia")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Separate departmental pack for:\n"
        "Kitchen staff · Restaurant / FOH staff · Building & Ground staff\n"
        "Anyone entering the kitchen, fridge, freezer, dry store or portable cabin"
    )
    set_run(run, size=11, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(14)
    run = meta.add_run(
        "The Vedanta Kitchen & Retreat Centre\n"
        "Branston Hall, Lincoln Road, Branston, Lincoln LN4 1PD\n"
        "Website: https://thevedanta.org/\n\n"
        "Assessor: Shyam Prasad\n"
        "Date: 27 June 2026\n"
        "Review due: 27 June 2027\n\n"
        "This pack is SEPARATE from the equipment risk assessment brochure.\n"
        "Issue a copy to each department whose staff enter the kitchen."
    )
    set_run(run, size=10, color=DARK)
    page_break(doc)

    # Contents
    add_heading(doc, "Contents", 1)
    for i, name in enumerate(
        [
            "Who this pack is for",
            "Mandatory PPE before entering the kitchen",
            "Restaurant & Front of House — kitchen access steps",
            "Moving safely around the kitchen",
            "Clean as you go — leave the kitchen professional",
            "Hot gastronorm trays, buffet & service collection",
            "Fridge, freezer, dry store & portable cabin",
            "Building / maintenance / other visiting staff",
            "Kitchen staff reminders",
            "Risk assessment table",
            "Emergency & reporting",
            "Department acknowledgement sign-off",
        ],
        1,
    ):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  {name}")
        set_run(run, size=11, color=DARK)
    page_break(doc)

    # 1 Who
    add_heading(doc, "1. Who this pack is for", 1)
    info_table(
        doc,
        [
            ("Workplace", "The Vedanta Kitchen & Retreat Centre — Branston Hall, LN4 1PD"),
            ("Company", "The Vedanta Way Limited"),
            ("Assessor", "Shyam Prasad"),
            ("Assessment date", "27 June 2026"),
            ("Review date", "27 June 2027, or sooner after an incident / near miss / layout change"),
            (
                "People covered",
                "Kitchen team; restaurant / FOH / service staff; housekeeping; maintenance; "
                "office / building / ground staff; contractors; any visitor entering food areas",
            ),
            (
                "Areas covered",
                "Main kitchen, pass/service area, fridge, freezer, dry store, portable cabin, "
                "corridors linked to kitchen, buffet / hot-holding collection points",
            ),
        ],
    )
    add_body(
        doc,
        "Anyone who is not working a full kitchen shift still creates risk when they enter the kitchen "
        "to collect food, check stock, clean, repair, or walk through. This pack sets one clear standard "
        "for all departments.",
    )
    page_break(doc)

    # 2 PPE
    add_heading(doc, "2. Mandatory PPE before entering the kitchen", 1)
    add_picture(doc, "staff-ppe-apron-hairnet-hands.png")
    add_body(doc, "No person may enter the kitchen food area without the following:", bold=True)
    add_numbered(
        doc,
        [
            "Anti-slip / non-slip health & safety footwear — mandatory. No open sandals, flip-flops or smooth soles.",
            "Hair net — all hair must be fully covered. Long hair secured under the net.",
            "Shirt apron — clean shirt apron must be worn when entering / using the kitchen.",
            "Wash hands when you enter the kitchen, and wash hands regularly while working with food, trays or storage.",
            "Remove loose jewellery, dangling lanyards and scarves that can catch on equipment.",
            "Wash or sanitise hands again before touching ready-to-eat food, clean trays or utensils.",
        ],
    )
    add_body(
        doc,
        "Note: A chef’s tissue / neck tissue is optional at this site. The required standard is "
        "anti-slip shoes + hair net + shirt apron + hand washing on entry and regularly.",
        size=10,
    )
    add_do_dont_care(
        doc,
        [
            "Wear anti-slip shoes every time",
            "Cover hair fully with a hair net",
            "Wear a clean shirt apron",
            "Wash hands when entering the kitchen and regularly during use",
        ],
        [
            "Don’t enter in open shoes or smooth soles",
            "Don’t leave hair uncovered",
            "Don’t enter without a shirt apron",
            "Don’t touch food or trays with unwashed hands",
            "Don’t wear outdoor coats over food-contact work",
        ],
        [
            "Keep spare hair nets and clean shirt aprons near the kitchen entrance",
            "Replace damaged or dirty PPE immediately",
            "Managers enforce the PPE gate — no exceptions for “quick visits”",
            "No hair net / shirt apron / hand wash = no kitchen entry",
        ],
    )
    page_break(doc)

    # 3 FOH steps
    add_heading(doc, "3. Restaurant & Front of House — kitchen access steps", 1)
    add_picture(doc, "staff-foh-kitchen-steps.png")
    add_body(
        doc,
        "Restaurant and FOH staff often enter to collect plated food, buffet gastronorm trays, "
        "drinks support items, or stock. Follow these steps every time:",
    )
    add_numbered(
        doc,
        [
            "Put on anti-slip shoes, hair net and shirt apron, then wash hands before entering.",
            "Pause at the door — look and listen. If the line is busy, wait or ask the chef before crossing.",
            "Watch the floor for water, oil, food spill or wet-floor signs.",
            "Keep clear of chefs’ knife work and prep boards; never reach across a cutting board.",
            "Stay away from open oven doors, hot ranges and fryers — steam and oil splash burn quickly.",
            "Use only agreed routes to the pass, fridge, freezer, dry store or cabin.",
            "When collecting hot trays / buffet items, use dry heat-proof cloths (see next section).",
            "If you used any equipment, board or knife — clean as you go and put everything back (see Clean as you go section).",
            "Leave promptly, keep walkways clear, and report any spill, out-of-date food or unsafe practice.",
        ],
    )
    page_break(doc)

    # 4 Moving safely
    add_heading(doc, "4. Moving safely around the kitchen", 1)
    add_picture(doc, "staff-kitchen-hazards-walk.png")
    add_body(doc, "Kitchen hazards for anyone walking through:", bold=True)
    for item in [
        "Wet, greasy or oily floors — high slip risk near sinks, dishwash, fryers and pass.",
        "Chefs carrying knives, hot pans or trays — call “behind”, “hot”, “sharp” and give way.",
        "Ovens on — doors may open suddenly with steam and radiant heat.",
        "Fryers in use — oil splash, hot zone around the fryer, never rush past.",
        "Trailing cables, trolleys, crates and bin bags in corridors.",
        "Fridge / freezer doors swinging into walkways.",
        "Noise and distraction — do not use phones while moving through the cook line.",
    ]:
        add_bullet(doc, item)
    add_do_dont_care(
        doc,
        [
            "Walk, don’t run",
            "Keep left / use agreed routes",
            "Call out when passing (“behind”, “hot”)",
            "Wipe or report spills immediately",
        ],
        [
            "Don’t cut through the hot cook line unless required",
            "Don’t stop to chat in doorways or at the pass",
            "Don’t carry more than you can see over",
            "Don’t ignore wet-floor signs",
        ],
        [
            "Expect the unexpected — ovens, fryers and knives move fast",
            "If unsure, wait for a kitchen team member to guide you",
            "Near-misses must be reported the same day",
        ],
    )
    page_break(doc)

    # 5 Clean as you go
    add_heading(doc, "5. Clean as you go — leave the kitchen professional", 1)
    add_picture(doc, "staff-clean-as-you-go.png")
    add_body(
        doc,
        "If restaurant staff, house staff, FOH, kitchen team or anyone else uses the kitchen — "
        "even for a short task — they must clean as they go. The kitchen must be left professional, "
        "tidy and ready for the next person. Do not leave chopping boards, knives, peelers, pans, "
        "utensils or equipment behind dirty or out of place.",
    )
    add_body(doc, "Clean-as-you-go steps (everyone):", bold=True)
    add_numbered(
        doc,
        [
            "Finish your food task first, then clean immediately — do not walk away “for later”.",
            "Wash or sanitise knives, peelers and utensils you used; dry them.",
            "Wash chopping boards (correct colour board for the food type) and stand them to dry or return clean.",
            "Wipe the worktop, sink edge and any spills on the floor you caused.",
            "Return knives and peelers to the designated drawer; return other equipment to its correct place.",
            "Never leave a dirty board, knife, blender shaft, pan or tray on the side for someone else to clear.",
            "If you used a fridge, freezer, dry store or cabin — close doors, wipe drips, put stock back neatly.",
            "Check the area looks professional before you leave. If you are unsure how to clean a piece of equipment, ask the chef — do not abandon it dirty.",
        ],
    )
    add_do_dont_care(
        doc,
        [
            "Clean boards, knives and tools straight after use",
            "Put everything back in the correct place / drawer",
            "Wipe worktops and floor spills you made",
            "Leave the kitchen tidy for the next person",
        ],
        [
            "Don’t leave chopping boards or knives on the side",
            "Don’t leave dirty equipment for kitchen staff to clear",
            "Don’t walk away mid-task without cleaning",
            "Don’t put wet knives loose and jumbled in the drawer",
        ],
        [
            "Clean as you go is mandatory for every department using the kitchen",
            "A professional kitchen stays ready for service at all times",
            "Managers may stop kitchen access for staff who repeatedly leave mess behind",
        ],
    )
    page_break(doc)

    # 6 Hot trays
    add_heading(doc, "6. Hot gastronorm trays, buffet & service collection", 1)
    add_picture(doc, "staff-hot-trays.png")
    add_body(
        doc,
        "Restaurant staff collecting food for buffet / service often handle hot gastronorm (GN) trays. "
        "Bare hands on hot trays cause burns; rushing with hot trays on wet floors causes slips and scalds.",
    )
    add_numbered(
        doc,
        [
            "Confirm the tray is ready with the kitchen / pass before lifting.",
            "Use dry heat-proof cloths or oven gloves — never wet cloths (steam burns) and never bare hands.",
            "Check lids are secure; beware of hot condensate under lids.",
            "Lift with two hands; keep the tray level; do not stack unstable hot loads.",
            "Announce “hot coming through” on the route to the buffet / service point.",
            "Walk carefully — if the floor is wet or oily, ask for the route to be cleared or dried first.",
            "Set trays onto a stable stand / Bain-marie / buffet well — not on an unstable edge.",
            "If a tray is too heavy or too hot, stop and ask for help. Do not struggle alone.",
        ],
    )
    risk_table(
        doc,
        [
            (
                "Burns from hot GN trays / lids",
                "Restaurant / FOH / kitchen staff",
                "High",
                "Mandatory dry heat-proof cloths or gloves; two-hand lift; announce hot; training before unsupervised collection.",
                "Low",
            ),
            (
                "Slip/trip while carrying hot food",
                "Restaurant / FOH staff; nearby guests/staff",
                "High",
                "Anti-slip shoes; clear route; no rushing; wipe spills; don’t carry overloaded stacks.",
                "Low",
            ),
            (
                "Scald from hot liquids spilling from tray",
                "Carrier and people nearby",
                "High",
                "Level carry; secure lids; don’t overfill; keep guests clear of route.",
                "Low",
            ),
        ],
    )
    page_break(doc)

    # 6 Storage
    add_heading(doc, "7. Fridge, freezer, dry store & portable cabin", 1)
    add_picture(doc, "staff-storage-access.png")
    add_body(
        doc,
        "Staff entering cold storage, dry store or the portable cabin must protect food safety as well as personal safety.",
    )
    add_numbered(
        doc,
        [
            "Wear required PPE before entry (anti-slip shoes, hair net, shirt apron) and wash hands.",
            "Open doors carefully — watch for people on the other side and icy / wet floors.",
            "Do not leave fridge or freezer doors open longer than needed.",
            "Check use-by / best-before dates before taking or returning products.",
            "Do not use out-of-date, damaged, unlabelled or unclean packaging — report it.",
            "Keep raw and ready-to-eat foods separate; never place raw above ready-to-eat.",
            "Maintain cleanliness: wipe drips, report dirty shelves, pests, bad smells or broken seals.",
            "In the portable cabin: secure the door, keep stock tidy, switch lights off if required, report damage.",
            "Wash / sanitise hands after handling outer packaging before touching ready-to-eat food.",
        ],
    )
    add_do_dont_care(
        doc,
        [
            "Check expiry dates every time",
            "Close fridge/freezer doors promptly",
            "Keep storage clean and organised",
            "Report damaged / out-of-date stock",
        ],
        [
            "Don’t leave doors propped open",
            "Don’t store personal food with kitchen stock",
            "Don’t ignore spills or unclean shelves",
            "Don’t use unlabelled containers",
        ],
        [
            "Cold chain and allergen control protect guests",
            "If in doubt, ask the chef / kitchen manager before using stock",
            "Cabin and dry store are still food areas — same hygiene rules apply",
        ],
    )
    page_break(doc)

    # 7 Building staff
    add_heading(doc, "8. Building / maintenance / other visiting staff", 1)
    add_picture(doc, "staff-building-visit.png")
    add_body(
        doc,
        "Housekeeping, maintenance, contractors and other building staff may enter for repairs, deliveries, "
        "cleaning or inspections. Extra controls apply because they may be unfamiliar with kitchen flow.",
    )
    add_numbered(
        doc,
        [
            "Report to the kitchen manager / chef in charge before starting work.",
            "Wear anti-slip shoes, hair net and shirt apron; wash hands on entry (plus any extra PPE for the task).",
            "Stay only in the agreed work area; do not walk the cook line unless escorted.",
            "Do not touch food, prep surfaces, ovens, fryers or knives unless authorised and trained.",
            "Protect food from dust, tools and chemicals — cover or remove food first if needed.",
            "Clean as you go; remove tools and waste before leaving — never leave boards, knives or equipment dirty.",
            "If cooking is live, wait for a safe window — never distract a chef mid-task at hot equipment.",
        ],
    )
    page_break(doc)

    # 8 Kitchen staff
    add_heading(doc, "9. Kitchen staff reminders", 1)
    add_body(
        doc,
        "Kitchen team members remain responsible for warning visitors and keeping routes as safe as practical during service.",
    )
    for item in [
        "Call “hot”, “behind”, “sharp”, “oven opening”, “fryer” clearly.",
        "Do not leave knives on the edge of boards or pointing into walkways.",
        "Close oven doors when not actively loading/unloading.",
        "Mop oil/water promptly and use wet-floor signs.",
        "Help FOH lift awkward hot GN trays when asked.",
        "Challenge anyone entering without anti-slip shoes, hair net, shirt apron or without washing hands.",
        "Keep fridge/freezer/dry store/cabin organised so visiting staff can find items safely.",
        "Remind restaurant / house staff: clean as you go — do not leave boards, knives or equipment behind.",
        "Do not accept a hand-over of dirty prep areas from visiting departments — ask them to finish cleaning first.",
    ]:
        add_bullet(doc, item)
    page_break(doc)

    # 9 Full risk table
    add_heading(doc, "10. Risk assessment table — staff kitchen access", 1)
    risk_table(
        doc,
        [
            (
                "Entering kitchen without anti-slip shoes",
                "All visiting / kitchen staff",
                "High",
                "Mandatory anti-slip footwear policy; refuse entry if not worn; spare pairs available if site policy allows.",
                "Low",
            ),
            (
                "Hair / clothing contamination of food",
                "Guests / food handlers",
                "High",
                "Hair net mandatory; clean shirt apron; remove loose jewellery; wash hands on entry and regularly.",
                "Low",
            ),
            (
                "Slips on wet / oily floors",
                "All staff in kitchen",
                "High",
                "Anti-slip shoes; clean-as-you-go; wet-floor signs; no running; report spills.",
                "Low",
            ),
            (
                "Contact with knives / being cut while passing",
                "FOH / visitors / kitchen",
                "High",
                "Keep clear of prep; chefs announce sharp; knives stored safely after use; no reaching across boards.",
                "Low",
            ),
            (
                "Burns / scalds from ovens, steam, fryers",
                "Anyone near cook line",
                "High",
                "Stand clear of oven doors and fryers; chefs announce opening; visitors use agreed routes only.",
                "Low",
            ),
            (
                "Burns from hot buffet / GN trays",
                "Restaurant / FOH staff",
                "High",
                "Heat-proof cloths/gloves required; two-hand method; training; ask for help with heavy/hot loads.",
                "Low",
            ),
            (
                "Collision with chef carrying hot items",
                "All staff",
                "Medium",
                "Call-outs; keep-left; no stopping in pass; limit congestion during service.",
                "Low",
            ),
            (
                "Using out-of-date / unclean stock from fridge/freezer/store/cabin",
                "Guests (food safety) / staff",
                "High",
                "Check dates every time; report unclean or damaged stock; FIFO; manager checks.",
                "Low",
            ),
            (
                "Cold-store door / icy floor injury",
                "Anyone entering cold rooms",
                "Medium",
                "Anti-slip shoes; open carefully; keep floors dry; do not rush; close doors properly.",
                "Low",
            ),
            (
                "Untrained building staff distracting cooks / touching equipment",
                "Kitchen team / visitors",
                "High",
                "Report in first; escort if needed; agreed work window; no unauthorised equipment use.",
                "Low",
            ),
            (
                "Leaving dirty boards, knives or equipment behind after use",
                "All staff using kitchen; next users; guests (hygiene)",
                "High",
                "Mandatory clean as you go; wash and return tools; wipe surfaces; managers enforce professional leave-behind standard.",
                "Low",
            ),
        ],
    )
    page_break(doc)

    # 10 Emergency
    add_heading(doc, "11. Emergency & reporting", 1)
    for item in [
        "Burns/scalds: cool under lukewarm running water for at least 20 minutes; seek first aid; call 999 for serious burns.",
        "Cuts: apply pressure with a clean dressing; seek first aid; do not continue food handling with an uncovered wound.",
        "Slips/falls: do not move a seriously injured person; call first aid / 999 as needed; secure the area.",
        "Food safety concern (out-of-date, contamination, allergen doubt): stop use, isolate product, tell the kitchen manager immediately.",
        "All accidents and near misses must be recorded the same day.",
    ]:
        add_bullet(doc, item)
    page_break(doc)

    # 11 Sign-off
    add_heading(doc, "12. Department acknowledgement sign-off", 1)
    add_body(
        doc,
        "I confirm I have read and understood this Staff Kitchen Access Risk Assessment, "
        "including mandatory anti-slip shoes, hair net, shirt apron, hand washing on entry and regularly, "
        "safe movement, hot tray handling, fridge/freezer/dry store/cabin rules, and clean as you go "
        "(never leave boards, knives or equipment dirty). "
        "I will follow these controls whenever I enter or use the kitchen.",
    )
    add_heading(doc, "Department copies to issue", 2)
    for dept in [
        "Kitchen team",
        "Restaurant / Front of House / service team",
        "Housekeeping",
        "Maintenance / facilities",
        "Office / building management",
        "Any contractor working in or through the kitchen",
    ]:
        add_bullet(doc, dept)

    table = doc.add_table(rows=9, cols=4)
    for i, h in enumerate(["Name", "Department / role", "Signature", "Date"]):
        write_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=9, center=True, fill=ALT_HEADER)
    for r in range(1, 9):
        for c in range(4):
            write_cell(table.rows[r].cells[c], "", size=9, fill=LIGHT_ROW if r % 2 == 0 else None)

    doc.add_paragraph()
    info_table(
        doc,
        [
            ("Document owner", "Shyam Prasad"),
            ("Approved by", ""),
            ("Approval date", ""),
            ("Version", "1.4 Staff Kitchen Access — shirt apron, hair net, hand wash on entry"),
            ("Related document", "Vedanta_Kitchen_Safety_Brochure_Risk_Pack.docx (equipment only)"),
        ],
    )

    out = "/workspace/vedanta-risk-assessments/Vedanta_Staff_Kitchen_Access_Risk_Assessment.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
